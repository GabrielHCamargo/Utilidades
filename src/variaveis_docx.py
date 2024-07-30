import os
from docx import Document
import openpyxl
 
def ler_dados_excel(arquivo_excel):
    """
    Lê os dados de um arquivo Excel e retorna uma lista de dicionários.
    Cada dicionário representa uma linha do Excel, com as chaves sendo os cabeçalhos das colunas.
    """
    wb = openpyxl.load_workbook(arquivo_excel)
    sheet = wb.active
    dados_lista = []
    headers = [cell.value for cell in sheet[1]]
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if row:
            dados = {headers[i]: row[i] for i in range(len(row))}
            dados_lista.append(dados)
    return dados_lista

def substituir_variaveis(doc, dados):
    """
    Substitui as variáveis no documento Word pelos valores correspondentes.
    As variáveis no documento devem estar no formato {{nome_da_variavel}}.
    """
    # Substituição nos parágrafos
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            for chave, valor in dados.items():
                variavel = '{{' + str(chave) + '}}'
                if variavel in run.text:
                    run.text = run.text.replace(variavel, str(valor))
    
    # Substituição nas tabelas
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in dados.items():
                    variavel = '{{' + str(chave) + '}}'
                    if variavel in celula.text:
                        celula.text = celula.text.replace(variavel, str(valor))

print("""
Instruções para inserção de variáveis no documento:

As variáveis devem ser inseridas no formato {{nome da variavel}}, onde:

1. As chaves duplas {{ }} são obrigatórias.
2. O nome da variável deve estar entre as chaves, sem espaços antes ou depois.
3. O nome da variável pode conter letras, números e underscores.

Exemplos corretos:
- {{nome}}
- {{idade}}
- {{endereco completo}}
- {{valor_total}}

Exemplos incorretos:
- {{ nome }}  (espaços entre as chaves e o nome da variável)
- {nome}      (apenas uma chave de cada lado)

As variáveis serão substituídas pelos valores correspondentes do arquivo Excel.
Certifique-se de que os nomes das variáveis no documento correspondam exatamente
aos cabeçalhos das colunas no arquivo Excel.

Exemplo de arquivo Excel:
| nome | idade | endereco completo | valor_total |
|------|-------|-------------------|-------------|
| João | 30    | Rua A, 1          | 1000.00     |
| Maria| 25    | Rua B, 2          | 1500.50     |
""")


# Solicita ao usuário o caminho do arquivo modelo
arquivo_word = input("Digite o caminho do arquivo modelo Word: ")

# Solicita ao usuário o caminho do arquivo de variáveis (Excel)
arquivo_excel = input("Digite o caminho do arquivo Excel com as variáveis: ")

# Lê os dados do arquivo Excel
dados_lista = ler_dados_excel(arquivo_excel)

# Define o diretório de saída
output_dir = '.output'
os.makedirs(output_dir, exist_ok=True)

# Processa cada conjunto de dados e gera um novo documento
for i, dados in enumerate(dados_lista, 1):
    # Carrega o documento modelo
    doc = Document(arquivo_word)
    
    # Substitui as variáveis no documento
    substituir_variaveis(doc, dados)
    
    # Define o nome do arquivo de saída
    nome_arquivo = os.path.join(output_dir, f'documento_preenchido_{i}.docx')
    
    # Salva o novo documento
    try:
        doc.save(nome_arquivo)
        print(f"Documento {nome_arquivo} criado com sucesso!")
    except Exception as e:
        print(f"Erro ao salvar o documento {nome_arquivo}: {e}")

print("Todos os documentos foram criados e preenchidos!")
